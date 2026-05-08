import pdfplumber
import os
import base64
import fitz  # PyMuPDF
import io
from PIL import Image
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_UNDERLINE

# Playwright is optional - only needed for HTML screenshots
try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

# python-docx is optional - only needed for DOCX output
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def parse_fontname(fontname):
    """Extract actual font name from pdfplumber's fontname and map to standard fonts"""
    if '+' in fontname:
        name = fontname.split('+')[1]
    else:
        name = fontname
    
    # Map common PDF font names to standard web fonts
    font_map = {
        'ArialMT': 'Arial, sans-serif',
        'Arial-BoldMT': 'Arial, sans-serif',
        'TimesNewRomanPSMT': 'Times New Roman, serif',
        'TimesNewRomanPS-BoldMT': 'Times New Roman, serif',
        'CourierNewPSMT': 'Courier New, monospace',
    }
    return font_map.get(name, f'{name}, sans-serif')

def extract_images_pymupdf(pdf_path):
    """Extract images from PDF using PyMuPDF, returns dict mapping (page_num, xref) to image data"""
    images = {}
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            if base_image:
                images[(page_num, xref)] = {
                    'data': base_image['image'],
                    'ext': base_image['ext'],
                    'width': base_image['width'],
                    'height': base_image['height']
                }
    doc.close()
    return images

def render_page_as_image(pdf_path, page_num, dpi=150):
    """Render entire PDF page as image"""
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    
    # Render to RGB (no alpha) - white background
    pix = page.get_pixmap(matrix=mat, alpha=False)
    
    # Convert to PNG
    img_data = pix.tobytes("png")
    doc.close()
    return base64.b64encode(img_data).decode('utf-8'), pix.width, pix.height

def pdf_to_html(pdf_path, html_path):
    html_parts = [
        '<!DOCTYPE html>',
        '<html><head>',
        '<meta charset="utf-8">',
        '<style>',
        '* { margin: 0; padding: 0; box-sizing: border-box; }',
        'body { background: #f0f0f0; padding: 20px; font-family: sans-serif; }',
        '.pdf-page { position: relative; margin: 0 auto 20px; background: white; box-shadow: 0 0 10px rgba(0,0,0,0.3); }',
        'span { position: absolute; white-space: pre; line-height: 1; margin: 0; padding: 0; font-style: normal; font-weight: normal; }',
        'div.rect { position: absolute; }',
        'div.line { position: absolute; z-index: 10; }',
        'img { position: absolute; max-width: none; }',
        'table { border-collapse: collapse; }',
        'td { border: 1px solid #000; padding: 0; }',
        '</style></head><body>'
    ]
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_width = page.width
            page_height = page.height
            html_parts.append(f'<div class="pdf-page" style="width:{page_width}pt;height:{page_height}pt;">')
            
            # Check if page has extractable text
            has_text = page.chars and len(page.chars) > 0
            
            # For scanned PDFs (no text), render page as image FIRST (background layer)
            # Skip embedded image extraction since page is already rendered as complete image
            if not has_text:
                img_b64, img_w, img_h = render_page_as_image(pdf_path, page_num, dpi=150)
                html_parts.append(f'<img src="data:image/png;base64,{img_b64}" style="left:0;top:0;width:{page_width}pt;height:{page_height}pt;" />')
            else:
                # For PDFs with text, extract embedded images using PyMuPDF
                try:
                    doc = fitz.open(pdf_path)
                    image_list = doc.get_page_images(page_num, full=True)
                    for img_idx, img_info in enumerate(image_list):
                        try:
                            xref = img_info[0]
                            img_width = img_info[2]
                            img_height = img_info[3]

                            base_image = doc.extract_image(xref)
                            if not base_image:
                                continue

                            img_data = base_image['image']
                            img_ext = base_image['ext'].lower()

                            try:
                                img = Image.open(io.BytesIO(img_data))

                                # Check if image has transparency
                                has_transparency = img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info)

                                # If PNG format - keep as PNG (supports transparency in browser)
                                if img_ext == 'png':
                                    buf = io.BytesIO()
                                    img.save(buf, format='PNG')
                                    img_data = buf.getvalue()
                                # If has transparency (RGBA/LA/P) but not PNG - paste on white and convert to JPEG
                                elif has_transparency:
                                    background = Image.new('RGB', img.size, (255, 255, 255))
                                    if img.mode == 'RGBA':
                                        background.paste(img, mask=img.split()[3])
                                    elif img.mode == 'LA':
                                        background.paste(img, mask=img.split()[1])
                                    elif img.mode == 'P':
                                        img = img.convert('RGBA')
                                        background.paste(img, mask=img.split()[3])
                                    img = background
                                    img_ext = 'jpeg'
                                    buf = io.BytesIO()
                                    img.save(buf, format='JPEG', quality=95)
                                    img_data = buf.getvalue()
                                else:
                                    # No transparency, convert to JPEG
                                    img_ext = 'jpeg'
                                    buf = io.BytesIO()
                                    img.save(buf, format='JPEG', quality=95)
                                    img_data = buf.getvalue()
                            except Exception as e:
                                pass

                            mime = f'image/{img_ext}'
                            img_b64 = base64.b64encode(img_data).decode('utf-8')

                            if img_idx < len(page.images):
                                img_dict = page.images[img_idx]
                                x0 = img_dict.get('x0', 0)
                                y0 = img_dict.get('top', 0)
                                w = img_dict.get('width', base_image['width'])
                                h = img_dict.get('height', base_image['height'])
                            else:
                                x0, y0, w, h = 0, 0, base_image['width'], base_image['height']
                            
                            html_parts.append(f'<img src="data:{mime};base64,{img_b64}" style="left:{x0}pt;top:{y0}pt;width:{w}pt;height:{h}pt;" />')
                        except Exception as e:
                            print(f'PyMuPDF image error p{page_num+1}-{img_idx}: {e}')
                    doc.close()
                except Exception as e:
                    print(f'PyMuPDF error p{page_num+1}: {e}')
            
            # For scanned PDFs (no text), skip extracting rects/lines as they are false positives from the image
            if not has_text:
                # Skip rect and line extraction for scanned PDFs - page is already rendered as image
                pass
            else:
                # Extract rectangles (only for PDFs with text)
                for rect in page.rects:
                    try:
                        x0 = rect['x0']
                        y0 = rect['top']
                        w = rect['width']
                        h = rect['height']
                        
                        sc = rect.get('stroke_color')
                        fc = rect.get('fill_color')
                        lw = rect.get('line_width', 1)
                        
                        stroke = '#000000'
                        if sc and isinstance(sc, (tuple, list)):
                            stroke = '#{:02x}{:02x}{:02x}'.format(*[int(c*255) for c in sc])
                        
                        fill = 'transparent'
                        if fc and isinstance(fc, (tuple, list)):
                            fill = '#{:02x}{:02x}{:02x}'.format(*[int(c*255) for c in fc])
                        
                        html_parts.append(f'<div class="rect" style="left:{x0}pt;top:{y0}pt;width:{w}pt;height:{h}pt;border:{lw}pt solid {stroke};background:{fill};"></div>')
                    except Exception as e:
                        print(f'Rect error p{page_num+1}: {e}')
                
                # Extract lines (horizontal/vertical dividers) - only for PDFs with text
                for line in page.lines:
                    try:
                        x0 = line['x0']
                        x1 = line['x1']
                        y0 = line['top']
                        y1 = line['bottom']
                        raw_lw = line.get('linewidth', 1)
                        
                        sc = line.get('stroking_color')
                        stroke = '#000000'
                        if sc and isinstance(sc, (tuple, list)):
                            stroke = '#{:02x}{:02x}{:02x}'.format(*[int(c*255) for c in sc])
                            brightness = sum([int(c*255) for c in sc]) / 3 / 255
                        else:
                            brightness = 0
                        
                        dash = line.get('dash')
                        is_dashed = dash and dash[0] and len(dash[0]) > 1
                        
                        width = x1 - x0
                        height = y1 - y0
                        
                        if brightness > 0.8:
                            lw = 1
                        elif brightness > 0.5:
                            lw = max(1, raw_lw / 10)
                        else:
                            lw = min(max(1, raw_lw / 5), 3)
                        
                        if height == 0 and width > 0:
                            html_parts.append(f'<div class="line" style="left:{x0}pt;top:{y0}pt;width:{width}pt;height:{lw}pt;background:{stroke};"></div>')
                        elif width == 0 and height > 0:
                            html_parts.append(f'<div class="line" style="left:{x0}pt;top:{y0}pt;width:{lw}pt;height:{height}pt;background:{stroke};"></div>')
                    except Exception as e:
                        print(f'Line error p{page_num+1}: {e}')
            
            # Skip words extraction for scanned PDFs
            if not has_text:
                html_parts.append('</div>')
                continue
            
            # Extract and render tables (disabled - causing rendering issues)
            # try:
            #     tables = page.debug_tablefinder().tables
            #     for table in tables:
            #         # Skip if table is not a Table object
            #         if not hasattr(table, 'bbox'):
            #             continue
            #         
            #         table_bbox = table.bbox  # (x0, top, x1, bottom)
            #         if not isinstance(table_bbox, (tuple, list)) or len(table_bbox) != 4:
            #             continue
            #         
            #         t_x0, t_top, t_x1, t_bottom = table_bbox
            #         t_width = t_x1 - t_x0
            #         t_height = t_bottom - t_top
            #         
            #         # Render table
            #         html_parts.append(f'<table style="position:absolute;left:{t_x0}pt;top:{t_top}pt;width:{t_width}pt;height:{t_height}pt;border-collapse:collapse;">')
            #         
            #         # Render cells
            #         if not hasattr(table, 'cells'):
            #             html_parts.append('</table>')
            #             continue
            #         
            #         cells = table.cells
            #         for row in cells:
            #             html_parts.append('<tr>')
            #             for cell in row:
            #                 if not cell or not hasattr(cell, 'bbox'):
            #                     html_parts.append('<td></td>')
            #                     continue
            #                 
            #                 cell_bbox = cell.bbox
            #                 if not isinstance(cell_bbox, (tuple, list)) or len(cell_bbox) != 4:
            #                     html_parts.append('<td></td>')
            #                     continue
            #                 
            #                 c_x0, c_top, c_x1, c_bottom = cell_bbox
            #                 c_width = c_x1 - c_x0
            #                 c_height = c_bottom - c_top
            #                 
            #                 # Get cell text
            #                 try:
            #                     cell_text = page.crop(cell_bbox).extract_text()
            #                     cell_text = cell_text.strip() if cell_text else ''
            #                 except:
            #                     cell_text = ''
            #                 
            #                 # Default styles
            #                 border_style = '1pt solid #000000'
            #                 bg_color = 'transparent'
            #                 
            #                 html_parts.append(f'<td style="width:{c_width}pt;height:{c_height}pt;border:{border_style};background:{bg_color};padding:0;">{cell_text}</td>')
            #             html_parts.append('</tr>')
            #         html_parts.append('</table>')
            # except Exception as e:
            #     print(f'Table error p{page_num+1}: {e}')
            
            # Extract and render each character individually to preserve all spaces
            chars = page.chars
            if chars:
                prev = None
                for c in chars:
                    text = c['text']
                    x0 = c['x0']
                    top = c['top']
                    size = c.get('size', 12)
                    
                    # Add space if there's a significant gap between characters on same line
                    if prev and abs(top - prev['top']) < 2:
                        gap = x0 - prev['x1']
                        if gap > size * 0.3:  # More than 30% of font size
                            style = f"left:{prev['x1']}pt;top:{prev['top']}pt;width:{gap}pt;font-size:{size}pt;"
                            html_parts.append(f'<span style="{style}"> </span>')
                    
                    # Render the character
                    font = parse_fontname(c.get('fontname', 'Arial, sans-serif'))
                    color = c.get('color', (0,0,0))
                    
                    if isinstance(color, (tuple, list)):
                        hex_color = '#{:02x}{:02x}{:02x}'.format(*[int(c*255) for c in color])
                    else:
                        hex_color = '#000000'
                    
                    style = f"left:{x0}pt;top:{top}pt;font-family:{font};font-size:{size}pt;color:{hex_color};"
                    html_parts.append(f'<span style="{style}">{text}</span>')
                    
                    prev = c
            
            html_parts.append('</div>')
    
    html_parts.append('</body></html>')
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_parts))

def pdf_to_images(pdf_path, out_dir, resolution=150):
    os.makedirs(out_dir, exist_ok=True)
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            img = page.to_image(resolution=resolution)
            path = os.path.join(out_dir, f'pdf_page_{i+1}.png')
            img.save(path, 'PNG')
            print(f'Saved PDF image: {path}')

def html_to_images(html_path, pdf_path, out_dir):
    """Convert HTML pages to images using Playwright"""
    if not PLAYWRIGHT_AVAILABLE:
        print('Error: Playwright is not installed.')
        print('To use HTML screenshot feature, install Playwright:')
        print('  pip install playwright')
        print('  playwright install chromium')
        return

    os.makedirs(out_dir, exist_ok=True)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(f'file://{os.path.abspath(html_path)}')
        page.wait_for_load_state('networkidle')

        pdf_pages = page.locator('.pdf-page')
        count = pdf_pages.count()

        for i in range(count):
            pdf_page_elem = pdf_pages.nth(i)
            bbox = pdf_page_elem.bounding_box()

            if bbox:
                width_px = int(bbox['width'])
                height_px = int(bbox['height'])

                page.set_viewport_size({'width': width_px + 100, 'height': height_px + 100})

                path = os.path.join(out_dir, f'html_page_{i+1}.png')
                pdf_page_elem.screenshot(path=path)
                print(f'Saved HTML image: {path}')

        browser.close()


def html_to_docx(html_path, docx_path):
    """Convert HTML to DOCX format"""
    if not DOCX_AVAILABLE:
        print('Error: python-docx is not installed.')
        print('To create DOCX files, install python-docx:')
        print('  pip install python-docx')
        return False

    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
    except Exception as e:
        print(f'Error reading HTML file: {e}')
        return False

    soup = BeautifulSoup(html_content, 'lxml')

    # Get page dimensions from style
    page_width = 595  # default A4 width in points
    page_height = 842  # default A4 height in points
    style = soup.find('style')
    if style:
        css = style.string or ''
        width_match = re.search(r'width:\s*(\d+)pt', css)
        height_match = re.search(r'height:\s*(\d+)pt', css)
        if width_match:
            page_width = int(width_match.group(1))
        if height_match:
            page_height = int(height_match.group(1))

    # Create DOCX document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(page_width / 72)
    section.page_height = Inches(page_height / 72)

    # Get all elements
    body = soup.find('body')
    if not body:
        print('Error: No body found in HTML')
        return False

    # Process each element in order (by top position)
    elements = []
    for elem in body.find_all(['div', 'p', 'img', 'table', 'hr']):
        style = elem.get('style', '')
        top_match = re.search(r'top:\s*([\d.]+)pt', style)
        left_match = re.search(r'left:\s*([\d.]+)pt', style)
        width_match = re.search(r'width:\s*([\d.]+)pt', style)
        height_match = re.search(r'height:\s*([\d.]+)pt', style)

        elem_data = {
            'element': elem,
            'top': float(top_match.group(1)) if top_match else 0,
            'left': float(left_match.group(1)) if left_match else 0,
            'width': float(width_match.group(1)) if width_match else page_width,
            'height': float(height_match.group(1)) if height_match else 0,
            'type': elem.name
        }
        elements.append(elem_data)

    # Sort by top position
    elements.sort(key=lambda x: x['top'])

    # Process each element
    for elem_data in elements:
        elem = elem_data['element']
        elem_type = elem_data['type']

        if elem_type == 'img':
            # Handle images
            src = elem.get('src', '')
            if src.startswith('data:image'):
                # Extract base64 image
                match = re.search(r'data:image/(\w+);base64,(.+)', src)
                if match:
                    img_data = base64.b64decode(match.group(2))
                    try:
                        img = Image.open(io.BytesIO(img_data))
                        # Save to temporary file for docx
                        temp_img_path = os.path.join(os.path.dirname(docx_path), 'temp_img.png')
                        img.save(temp_img_path)

                        # Calculate image size in inches
                        width_inch = elem_data['width'] / 72
                        height_inch = elem_data['height'] / 72

                        # Add to docx (max width 6 inches)
                        max_width = 6
                        if width_inch > max_width:
                            ratio = max_width / width_inch
                            width_inch = max_width
                            height_inch = height_inch * ratio

                        doc.add_picture(temp_img_path, width=Inches(width_inch), height=Inches(height_inch))

                        # Clean up temp file
                        try:
                            os.remove(temp_img_path)
                        except:
                            pass
                    except Exception as e:
                        print(f'Warning: Could not add image: {e}')

        elif elem_type == 'table':
            # Handle tables
            rows = elem.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])) if rows else 0)
                table.style = 'Table Grid'

                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for c_idx, cell in enumerate(cells):
                        if r_idx < len(table.rows) and c_idx < len(table.columns):
                            cell_text = cell.get_text(strip=True)
                            table.cell(r_idx, c_idx).text = cell_text

        elif elem_type == 'hr':
            # Horizontal line
            doc.add_paragraph('_' * 50)

        elif elem_type in ('div', 'p'):
            # Handle text content
            text = elem.get_text(strip=True)
            if text:
                # Check if it's a rectangle (border box with minimal text)
                style = elem.get('style', '')
                if 'border' in style or 'border-width' in style:
                    # This might be a rectangle, add as bordered paragraph
                    p = doc.add_paragraph(text)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()

                    # Parse inline styles
                    for span in elem.find_all('span'):
                        span_text = span.get_text()
                        if not span_text:
                            continue

                        # Get style
                        span_style = span.get('style', '')

                        # Font size
                        font_size = 12
                        size_match = re.search(r'font-size:\s*([\d.]+)pt', span_style)
                        if size_match:
                            font_size = float(size_match.group(1))

                        # Font color
                        color = None
                        color_match = re.search(r'color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', span_style)
                        if color_match:
                            color = RGBColor(
                                int(color_match.group(1)),
                                int(color_match.group(2)),
                                int(color_match.group(3))
                            )

                        # Font name
                        font_name = 'Calibri'
                        font_match = re.search(r'font-family:\s*([^;]+)', span_style)
                        if font_match:
                            font_name = parse_fontname(font_match.group(1).strip())

                        # Bold, Italic, Underline
                        is_bold = 'font-weight: bold' in span_style or 'font-weight: 700' in span_style
                        is_italic = 'font-style: italic' in span_style
                        is_underline = 'text-decoration: underline' in span_style

                        # Add run
                        run = p.add_run(span_text)
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        if color:
                            run.font.color.rgb = color
                        if is_bold:
                            run.font.bold = True
                        if is_italic:
                            run.font.italic = True
                        if is_underline:
                            run.font.underline = WD_UNDERLINE.SINGLE

                    # If no spans found, add plain text
                    if not p.runs:
                        run = p.add_run(text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)

    try:
        doc.save(docx_path)
        print(f'Saved DOCX: {docx_path}')
        return True
    except Exception as e:
        print(f'Error saving DOCX: {e}')
        return False


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Convert PDF to HTML/DOCX/DOC')
    parser.add_argument('pdf_path', help='Path to input PDF file')
    parser.add_argument('-o', '--output', default='output.html', help='Output file path')
    parser.add_argument('-f', '--format', default='html', choices=['html', 'docx', 'doc'], help='Output format')
    parser.add_argument('--screenshot', nargs='+', default=[], choices=['pdf', 'html', 'all'], help='Generate screenshots: pdf, html, all')
    args = parser.parse_args()

    pdf_path = args.pdf_path
    output_path = args.output
    output_format = args.format

    if not os.path.exists(pdf_path):
        print(f'Error: PDF file not found: {pdf_path}')
        return

    # Determine output path with correct extension
    if output_format == 'html' and not output_path.endswith('.html'):
        output_path = output_path.replace('.docx', '').replace('.doc', '') + '.html'
    elif output_format == 'docx' and not output_path.endswith('.docx'):
        output_path = output_path.replace('.html', '').replace('.doc', '') + '.docx'
    elif output_format == 'doc' and not output_path.endswith('.doc'):
        output_path = output_path.replace('.html', '').replace('.docx', '') + '.doc'

    print(f'Converting PDF to {output_format.upper()}: {pdf_path}')

    # Generate HTML first (intermediate format)
    html_path = output_path.replace('.docx', '.doc').replace('.doc', '') + '.html'
    pdf_to_html(pdf_path, html_path)

    # If output format is docx, convert HTML to DOCX
    if output_format == 'docx':
        success = html_to_docx(html_path, output_path)
        if not success:
            print('Failed to create DOCX, HTML file is available instead.')
            return

        # Clean up intermediate HTML file
        if html_path != output_path:
            try:
                os.remove(html_path)
            except:
                pass

    screenshot_opts = args.screenshot
    if 'pdf' in screenshot_opts or 'all' in screenshot_opts:
        print('Converting PDF pages to images...')
        pdf_to_images(pdf_path, 'pdf_images')

    if 'html' in screenshot_opts or 'all' in screenshot_opts:
        print('Converting HTML to images...')
        html_to_images(html_path, pdf_path, 'html_images')

    print(f'Done! Output: {output_path}')
    if 'pdf' in screenshot_opts or 'all' in screenshot_opts:
        print('Check pdf_images/ folder for PDF screenshots.')
    if 'html' in screenshot_opts or 'all' in screenshot_opts:
        print('Check html_images/ folder for HTML screenshots.')

if __name__ == '__main__':
    main()
